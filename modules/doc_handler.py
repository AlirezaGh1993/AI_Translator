# modules/doc_handler.py

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
import math

# --- Constants ---
# Max characters to send per API call. Adjust based on model context window.
MAX_CHUNK_SIZE = 4000 

def is_rtl(lang):
    """Check if a language is Right-to-Left."""
    # A simple check, can be expanded
    return lang.lower() in ['fa', 'ar', 'he', 'ur']

def chunk_text(text, max_size):
    """Splits text into chunks by paragraphs without breaking them."""
    paragraphs = text.split('\n')
    chunks = []
    current_chunk = ""
    for p in paragraphs:
        if len(current_chunk) + len(p) + 1 > max_size:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = p
        else:
            if current_chunk:
                current_chunk += "\n" + p
            else:
                current_chunk = p
    if current_chunk:
        chunks.append(current_chunk)
    return chunks

def translate_docx(filepath, output_path, translate_func, api_key, style, target_lang, proxy_url, update_progress):
    """Translates a .docx file while preserving styles and images."""
    document = Document(filepath)
    total_paragraphs = len(document.paragraphs)
    processed_count = 0

    for i, para in enumerate(document.paragraphs):
        if para.text.strip(): # Only translate non-empty paragraphs
            original_text = para.text
            # Translate paragraph by paragraph
            translated_text = translate_func(api_key, original_text, style, target_lang, proxy_url)
            
            # Clear existing paragraph content but keep formatting
            para.clear()
            # Add translated text with original styling (runs)
            # This is a simplified approach; for complex formatting, run-level translation is needed
            run = para.add_run(translated_text)
            
            # Apply RTL settings if needed
            if is_rtl(target_lang):
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.paragraph_format.right_to_left = True

        processed_count += 1
        progress = math.ceil((processed_count / total_paragraphs) * 100)
        update_progress(progress, f"Translating paragraph {processed_count}/{total_paragraphs}...")

    document.save(output_path)

def translate_txt(filepath, output_path, translate_func, api_key, style, target_lang, proxy_url, update_progress):
    """Translates a plain .txt file."""
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    
    text_chunks = chunk_text(text, MAX_CHUNK_SIZE)
    total_chunks = len(text_chunks)
    translated_chunks = []

    for i, chunk in enumerate(text_chunks):
        translated_chunk = translate_func(api_key, chunk, style, target_lang, proxy_url)
        translated_chunks.append(translated_chunk)
        progress = math.ceil(((i + 1) / total_chunks) * 100)
        update_progress(progress, f"Translating chunk {i+1}/{total_chunks}...")
        
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(translated_chunks))

def translate_pdf(filepath, output_path, translate_func, api_key, style, target_lang, proxy_url, update_progress):
    """Translates a .pdf file and saves the output as a .docx file."""
    doc = fitz.open(filepath)
    full_text = ""
    for page in doc:
        full_text += page.get_text("text") + "\n"
    doc.close()

    text_chunks = chunk_text(full_text, MAX_CHUNK_SIZE)
    total_chunks = len(text_chunks)
    translated_chunks = []

    for i, chunk in enumerate(text_chunks):
        translated_chunk = translate_func(api_key, chunk, style, target_lang, proxy_url)
        translated_chunks.append(translated_chunk)
        progress = math.ceil(((i + 1) / total_chunks) * 100)
        update_progress(progress, f"Translating chunk {i+1}/{total_chunks}...")

    # Create a new docx for the output
    output_doc = Document()
    for translated_text in translated_chunks:
        for para_text in translated_text.split('\n'):
            para = output_doc.add_paragraph(para_text)
            if is_rtl(target_lang):
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.paragraph_format.right_to_left = True
    
    # PDF output becomes DOCX because styling can't be preserved
    final_output_path = os.path.splitext(output_path)[0] + '_translated.docx'
    output_doc.save(final_output_path)
    return final_output_path
